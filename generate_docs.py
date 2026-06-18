import pandas as pd
from docx import Document

df = pd.read_excel("data.xlsx")

for _, row in df.iterrows():
    doc = Document("template.docx")

    # Paragraphs
    for para in doc.paragraphs:
        para.text = para.text.replace("{{NAME}}", str(row["NAME"]))
        para.text = para.text.replace("{{SERIAL}}", str(row["SERIAL"]))

    # Tables
    for table in doc.tables:
        for table_row in table.rows:
            for cell in table_row.cells:
                cell.text = cell.text.replace("{{NAME}}", str(row["NAME"]))
                cell.text = cell.text.replace("{{SERIAL}}", str(row["SERIAL"]))

    doc.save(f"{row['NAME']}.docx")